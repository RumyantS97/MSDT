from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Documents:
    """
    Класс для создания документов, таких как освобождение и благодарность.
    """
    exemption_doc: Document
    thanks_doc: Document
    _sample_dir: str
    _doc_style = 'Normal'
    _font_style = 'Times New Roman'
    _font_size = 14

    def __init__(self):
        self.date = 'date'
        self.event_name = 'event name'
        self.institute_name = 'institute name'
        self.director_name = 'director name'
        self.full_names = []
        self.signature_post = 'Post of person'
        self.signature_name = 'person'

    def set_data(self, date: str):
        """
        Устанавливает дату мероприятия.
        """
        self.date = date

    def set_event_name(self, event_name: str):
        """
        Устанавливает название мероприятия.
        """
        self.event_name = event_name

    def set_institute(
            self, institute_name: str, 
            director_post: str, director_name: str):
        """
        Устанавливает информацию об институте и директоре.
        """
        self.institute_name = institute_name
        self.director_name = director_name
        self.director_post = director_post

    def add_full_names(self, full_names: list[list[str]]):
        """
        Добавляет список студентов в таблицу.
        """
        self.full_names = full_names

    def add_fio(self, fio: list[str]):
        """
        Добавляет одного студента в таблицу.
        """
        self.full_names.append(fio)

    def set_signature(self, post: str, name: str):
        """
        Устанавливает должность и имя для подписи.
        """
        self.signature_post = post
        self.signature_name = name

    def _add_exemption_content(self):
        """
        Добавляет содержимое документа освобождения.
        Заполняет таблицу, добавляет подписи и заменяет
        текстовые плейсхолдеры на реальные значения.
        """
        # добавление института
        self.exemption_doc.paragraphs[1].add_run(self.institute_name)

        # добавление фио директора
        self.exemption_doc.paragraphs[2].add_run(self.director_name)

        # добавление названия мероприятия
        for i in self.exemption_doc.paragraphs:
            i.text = i.text.replace("event", self.event_name)

        # добавление даты мероприятия
        for j in self.exemption_doc.paragraphs:
            j.text = j.text.replace("date", self.date)

        # добавление подписи
        sign = self.exemption_doc.tables[1]
        p = sign.cell(0, 0).add_paragraph('\n' + self.signature_post)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p = sign.cell(0, 1).add_paragraph(self.signature_name)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Заполнение таблицы
        fio = self.exemption_doc.tables[0]
        i = 0
        for students_name, group in self.full_names:
            i += 1
            cells = fio.add_row().cells
            cells[0].text = str(i) + "."  # порядковый номер
            cells[1].text = students_name  # ФИО
            cells[2].text = group  # Группа

    def make_exemption(
            self, dir: str, 
            exemption_sample: str, file_name: str):
        """
        Создает документ освобождения на основе шаблона.
        """
        self.exemption_doc = Document(exemption_sample)
        style = self.exemption_doc.styles[self._doc_style]
        style.font.name = self._font_style
        style.font.size = Pt(self._font_size)
        self._add_exemption_content()
        self.exemption_doc.save(
            f"./{dir}/Освобождение {file_name}_{self.institute_name}.docx")

    def _add_gatitude_content(self):
        """
        Добавляет содержимое документа благодарности.
        Заполняет таблицу, добавляет подписи и заменяет
        текстовые плейсхолдеры на реальные значения.
        """
        run = self.thanks_doc.paragraphs[3].add_run(self.institute_name)
        run.bold = True
        fio = self.thanks_doc.tables[2]
        for i in self.thanks_doc.paragraphs:
            i.text = i.text.replace("event", self.event_name)

        # добавление даты мероприятия
        for j in self.thanks_doc.paragraphs:
            j.text = j.text.replace("date", self.date)
        i = 0

        # добавление подписи
        sign = self.thanks_doc.tables[3]
        p = sign.cell(0, 0).add_paragraph(self.signature_post)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p = sign.cell(0, 1).add_paragraph(self.signature_name)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for students_name, group in self.full_names:
            i += 1
            cells = fio.add_row().cells
            cells[0].text = str(i) + '.'  # порядковый номер
            cells[1].text = students_name  # фио студента
            cells[2].text = group  # группа студента

    def make_thanks(
            self, dir: str, 
            sample_thanks: str, file_name: str):
        """
        Создает документ благодарности на основе шаблона.
        """
        self.thanks_doc = Document(sample_thanks)
        style = self.thanks_doc.styles[self._doc_style]
        style.font.name = self._font_style
        style.font.size = Pt(self._font_size)
        self._add_gatitude_content()
        self.thanks_doc.save(f"./{dir}/Благодарность_{file_name}_{self.institute_name}.docx")
